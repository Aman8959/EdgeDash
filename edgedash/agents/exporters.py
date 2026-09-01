"""Resume Exporter - Export resume to DOCX format."""

from typing import Optional
from edgedash.models.resume import ResumeVersion


class DocxExporter:
    """Export resume to DOCX format using python-docx."""
    
    @staticmethod
    def export_to_docx(resume: ResumeVersion, output_path: str) -> bool:
        """Export resume to DOCX file.
        
        Args:
            resume: Resume version to export
            output_path: Path where to save DOCX file
        
        Returns:
            True if successful, False otherwise
        """
        try:
            from docx import Document
            from docx.shared import Pt, Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise ImportError("python-docx not installed. Run: pip install python-docx")
        
        try:
            # Create document
            doc = Document()
            
            # Set document margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)
                section.left_margin = Inches(0.75)
                section.right_margin = Inches(0.75)
            
            # Add header with name
            header = doc.add_paragraph()
            header.alignment = WD_ALIGN_PARAGRAPH.CENTER
            name_run = header.add_run(resume.full_name)
            name_run.font.size = Pt(16)
            name_run.font.bold = True
            
            # Add contact info
            contact = doc.add_paragraph()
            contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_run = contact.add_run(resume.contact_info)
            contact_run.font.size = Pt(10)
            
            # Add blank line
            doc.add_paragraph()
            
            # Professional Summary section
            if resume.professional_summary:
                summary_heading = doc.add_paragraph()
                summary_heading_run = summary_heading.add_run("PROFESSIONAL SUMMARY")
                summary_heading_run.font.bold = True
                summary_heading_run.font.size = Pt(11)
                
                summary_text = doc.add_paragraph(resume.professional_summary)
                summary_text.paragraph_format.space_before = Pt(0)
                summary_text.paragraph_format.space_after = Pt(6)
                
                doc.add_paragraph()
            
            # Skills section
            if resume.skills_section:
                skills_heading = doc.add_paragraph()
                skills_heading_run = skills_heading.add_run("SKILLS")
                skills_heading_run.font.bold = True
                skills_heading_run.font.size = Pt(11)
                
                # Add skills in 2-3 columns using table
                table = doc.add_table(rows=1, cols=2)
                table.autofit = False
                
                # Add skills to table cells
                skills = resume.skills_section
                for i in range(0, len(skills), 2):
                    if i + 1 < len(skills):
                        row = table.add_row()
                        row.cells[0].text = skills[i]
                        row.cells[1].text = skills[i + 1]
                    else:
                        row = table.add_row()
                        row.cells[0].text = skills[i]
                
                doc.add_paragraph()
            
            # Experience section
            if resume.experience_section:
                exp_heading = doc.add_paragraph()
                exp_heading_run = exp_heading.add_run("WORK EXPERIENCE")
                exp_heading_run.font.bold = True
                exp_heading_run.font.size = Pt(11)
                
                # Parse experience lines and format
                for line in resume.experience_section:
                    if not line.strip():
                        continue
                    
                    if '|' in line and not line.startswith('  '):
                        # Header line
                        exp_entry = doc.add_paragraph()
                        exp_entry_run = exp_entry.add_run(line.strip())
                        exp_entry_run.font.bold = True
                        exp_entry_run.font.size = Pt(10)
                    elif line.startswith('  •'):
                        # Bullet point
                        bullet = doc.add_paragraph(line.strip()[2:], style='List Bullet')
                        bullet.paragraph_format.space_before = Pt(0)
                        bullet.paragraph_format.space_after = Pt(3)
                    elif line.startswith('  '):
                        # Regular indented text
                        desc = doc.add_paragraph(line.strip())
                        desc.paragraph_format.space_before = Pt(0)
                        desc.paragraph_format.space_after = Pt(3)
                
                doc.add_paragraph()
            
            # Projects section
            if resume.projects_section:
                proj_heading = doc.add_paragraph()
                proj_heading_run = proj_heading.add_run("PROJECTS")
                proj_heading_run.font.bold = True
                proj_heading_run.font.size = Pt(11)
                
                for line in resume.projects_section:
                    if not line.strip():
                        continue
                    
                    if 'Tech:' in line or 'Link:' in line or line.startswith('  '):
                        # Sub-item
                        sub_item = doc.add_paragraph(line.strip())
                        sub_item.paragraph_format.space_before = Pt(0)
                        sub_item.paragraph_format.space_after = Pt(2)
                    else:
                        # Project header
                        proj_entry = doc.add_paragraph()
                        proj_entry_run = proj_entry.add_run(line.strip())
                        proj_entry_run.font.bold = True
                        proj_entry_run.font.size = Pt(10)
                
                doc.add_paragraph()
            
            # Education section
            if resume.education_section:
                edu_heading = doc.add_paragraph()
                edu_heading_run = edu_heading.add_run("EDUCATION")
                edu_heading_run.font.bold = True
                edu_heading_run.font.size = Pt(11)
                
                for edu in resume.education_section:
                    edu_entry = doc.add_paragraph(edu)
                    edu_entry.paragraph_format.space_before = Pt(0)
                    edu_entry.paragraph_format.space_after = Pt(6)
                
                doc.add_paragraph()
            
            # Certifications section
            if resume.certifications_section:
                cert_heading = doc.add_paragraph()
                cert_heading_run = cert_heading.add_run("CERTIFICATIONS")
                cert_heading_run.font.bold = True
                cert_heading_run.font.size = Pt(11)
                
                for cert in resume.certifications_section:
                    cert_entry = doc.add_paragraph(cert)
                    cert_entry.paragraph_format.space_before = Pt(0)
                    cert_entry.paragraph_format.space_after = Pt(3)
                
                doc.add_paragraph()
            
            # Achievements section
            if resume.achievements_section:
                ach_heading = doc.add_paragraph()
                ach_heading_run = ach_heading.add_run("ACHIEVEMENTS")
                ach_heading_run.font.bold = True
                ach_heading_run.font.size = Pt(11)
                
                for achievement in resume.achievements_section:
                    ach = doc.add_paragraph(achievement, style='List Bullet')
                    ach.paragraph_format.space_before = Pt(0)
                    ach.paragraph_format.space_after = Pt(3)
            
            # Save document
            doc.save(output_path)
            return True
        
        except Exception as e:
            raise ValueError(f"Failed to export resume to DOCX: {str(e)}")


class PlainTextExporter:
    """Export resume to plain text format."""
    
    @staticmethod
    def export_to_txt(resume: ResumeVersion, output_path: str) -> bool:
        """Export resume to plain text file.
        
        Args:
            resume: Resume version to export
            output_path: Path where to save text file
        
        Returns:
            True if successful, False otherwise
        """
        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(resume.content_text)
            return True
        except Exception as e:
            raise ValueError(f"Failed to export resume to text: {str(e)}")


class HTMLExporter:
    """Export resume to HTML format for web viewing."""
    
    @staticmethod
    def export_to_html(resume: ResumeVersion, output_path: str) -> bool:
        """Export resume to HTML file.
        
        Args:
            resume: Resume version to export
            output_path: Path where to save HTML file
        
        Returns:
            True if successful, False otherwise
        """
        try:
            html_content = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{resume.full_name} - {resume.target_role}</title>
    <style>
        body {{
            font-family: Arial, sans-serif;
            line-height: 1.6;
            max-width: 900px;
            margin: 0 auto;
            padding: 20px;
            color: #333;
        }}
        .header {{
            text-align: center;
            margin-bottom: 20px;
            border-bottom: 2px solid #333;
            padding-bottom: 10px;
        }}
        .name {{
            font-size: 24px;
            font-weight: bold;
        }}
        .contact {{
            font-size: 12px;
            color: #666;
        }}
        .section-title {{
            font-size: 14px;
            font-weight: bold;
            background-color: #f0f0f0;
            padding: 8px;
            margin-top: 15px;
            margin-bottom: 10px;
        }}
        .entry-header {{
            font-weight: bold;
            margin-top: 10px;
        }}
        .entry-detail {{
            margin-left: 20px;
            font-size: 13px;
        }}
        .skills-grid {{
            display: grid;
            grid-template-columns: 1fr 1fr;
            gap: 10px;
        }}
        .skill-item {{
            font-size: 13px;
        }}
        ul {{
            margin: 5px 0;
            padding-left: 20px;
        }}
        li {{
            font-size: 13px;
            margin: 3px 0;
        }}
    </style>
</head>
<body>
    <div class="header">
        <div class="name">{resume.full_name}</div>
        <div class="contact">{resume.contact_info}</div>
    </div>
    
    <div class="section-title">PROFESSIONAL SUMMARY</div>
    <div class="entry-detail">{resume.professional_summary}</div>
    
    <div class="section-title">SKILLS</div>
    <div class="skills-grid">
"""
            
            # Add skills
            for skill in resume.skills_section:
                html_content += f'        <div class="skill-item">{skill}</div>\n'
            
            html_content += """    </div>
    
    <div class="section-title">WORK EXPERIENCE</div>
"""
            
            # Add experience
            for line in resume.experience_section:
                if '|' in line and not line.startswith('  '):
                    html_content += f'    <div class="entry-header">{line.strip()}</div>\n'
                elif line.startswith('  •'):
                    html_content += f'    <div class="entry-detail"><ul><li>{line.strip()[2:]}</li></ul></div>\n'
                elif line.startswith('  '):
                    html_content += f'    <div class="entry-detail">{line.strip()}</div>\n'
            
            html_content += """
    <div class="section-title">EDUCATION</div>
"""
            
            # Add education
            for edu in resume.education_section:
                html_content += f'    <div class="entry-detail">{edu}</div>\n'
            
            # Add projects if available
            if resume.projects_section:
                html_content += """
    <div class="section-title">PROJECTS</div>
"""
                for line in resume.projects_section:
                    if line.strip() and not line.startswith('  '):
                        html_content += f'    <div class="entry-header">{line.strip()}</div>\n'
                    elif line.strip():
                        html_content += f'    <div class="entry-detail">{line.strip()}</div>\n'
            
            html_content += """
</body>
</html>"""
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            return True
        
        except Exception as e:
            raise ValueError(f"Failed to export resume to HTML: {str(e)}")
